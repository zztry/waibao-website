# import win32com.client as win32 #docx2pdf
from docx import Document
import docx.shared
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from pydocx import PyDocX
import os
import processing


def word_output(score, problems: list[tuple[int, str, str, list[str], float]], suggestion: str, problem_analysis: str):
    doc = Document()
    head = doc.add_paragraph('用户体验评分报告')
    run1 = head.runs[0]
    run1.font.name = '楷体'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run1.font.size = Pt(32)
    run1.font.color.rgb = RGBColor(54, 95, 145)
    head.alignment = 1

    doc.add_paragraph("")
    para0 = doc.add_paragraph("一、用户体验评分")
    run = para0.runs[0]
    run.font.name = '楷体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(54, 95, 145)

    scores = '用户体验评分为： ' + score + '分。'
    paragraph = doc.add_paragraph(scores)
    run2 = paragraph.runs[0]
    run2.font.name = '楷体'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run2.font.size = Pt(20)
    paragraph.paragraph_format.first_line_indent = docx.shared.Inches(0.5)

    doc.add_paragraph("")
    para1 = doc.add_paragraph("二、存在的问题")
    run3 = para1.runs[0]
    run3.font.name = '楷体'
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run3.font.size = Pt(24)
    run3.font.color.rgb = RGBColor(54, 95, 145)

    problem_analysis_split = problem_analysis.split("\n\t")
    for part in problem_analysis_split:
        if len(part) != 0:
            para4 = doc.add_paragraph(part)
            run5 = para4.runs[0]
            run5.font.name = '楷体'
            run5._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            run5.font.size = Pt(16)
            para4.paragraph_format.first_line_indent = docx.shared.Inches(0.5)

    # 三、开发者建议
    doc.add_paragraph("")
    para2 = doc.add_paragraph("三、开发者建议")
    run3 = para2.runs[0]
    run3.font.name = '楷体'
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run3.font.size = Pt(24)
    run3.font.color.rgb = RGBColor(54, 95, 145)

    suggestions = suggestion.split("\n\t")
    for part in suggestions:
        if len(part) != 0:
            para3 = doc.add_paragraph(part)
            run4 = para3.runs[0]
            run4.font.name = '楷体'
            run4._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            run4.font.size = Pt(16)
            para3.paragraph_format.first_line_indent = docx.shared.Inches(0.5)

    output_name = 'output_report.docx'
    doc.save(output_name)
    #print("111")
    return doc


# def convert2pdf(input_path, output_path):
#     word_app = win32.gencache.EnsureDispatch('word.Application')
#     word_app.Visible = False
#     try:
#         doc = word_app.Documents.Open(input_path)
#         doc.SaveAs(output_path, FileFormat=17)
#         doc.Close()
#         return True
#     except Exception as e:
#         print("转换失败：" + str(e))
#         return False
#     finally:
#         word_app.Quit()

# testpath: list[str] = ['/Users/qinhaonan/Desktop/NJU_FILES/服务外包/data_sets/log3-new.json',
#                       '/Users/qinhaonan/Desktop/NJU_FILES/服务外包/data_sets/log2-new.json']

input_path = os.getcwd() + '/output_docx.docx'
output_path = os.getcwd() + '/output_pdf'

# probs = processing.analyze_file(testpath)
# word_output('90', probs)
